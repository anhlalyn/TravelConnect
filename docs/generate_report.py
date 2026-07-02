import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.style import WD_STYLE_TYPE
import os
import shutil
from PIL import Image

# Paths
base_dir = r"d:\Deadline\ChuyenDe2\ChuyenDe"
original_docx = os.path.join(base_dir, "docs", "WordChuyenDe2.docx")
output_docx = os.path.join(base_dir, "docs", "BaoCaoChuyenDe1_PhanDinhLuyen.docx")
scratch_img_dir = r"C:\Users\anhlalyn\.gemini\antigravity-ide\brain\271da160-8c6f-4f31-9a4e-8d998b9cd656\scratch\downloaded_images"
extracted_img_dir = r"C:\Users\anhlalyn\.gemini\antigravity-ide\brain\271da160-8c6f-4f31-9a4e-8d998b9cd656\scratch\extracted_images"
dest_img_dir = os.path.join(base_dir, "docs", "images")

if not os.path.exists(dest_img_dir):
    os.makedirs(dest_img_dir)

# 1. Copy diagrams and logo to workspace images folder
diagram_files = {
    "usecase.png": "usecase.png",
    "dfd_context.png": "dfd_context.png",
    "dfd_level1.png": "dfd_level1.png",
    "erd.png": "erd.png",
    "userflow.png": "userflow.png",
    "flow_register.png": "flow_register.png",
    "flow_explore.png": "flow_explore.png",
    "flow_booking.png": "flow_booking.png",
    "flow_business.png": "flow_business.png",
    "flow_admin.png": "flow_admin.png",
    "flow_forgot_password.png": "flow_forgot_password.png",
    "s__lung_hot_ng_h_thng.png": "flow_system.png",
    "deployment.png": "deployment.png",
    "screenshot_homepage.png": "screenshot_homepage.png",
    "screenshot_dashboard.png": "screenshot_dashboard.png"
}

# Copy logo from extracted images (actual logo is img_6_image2.png)
logo_path = os.path.join(extracted_img_dir, "img_6_image2.png")
dest_logo_path = os.path.join(dest_img_dir, "logo.png")
if os.path.exists(logo_path):
    shutil.copy(logo_path, dest_logo_path)
    print("Logo copied to", dest_logo_path)
else:
    # try fallback if img_6_image2 doesn't exist
    fallback_path = os.path.join(extracted_img_dir, "img_0_image1.png")
    if os.path.exists(fallback_path):
        shutil.copy(fallback_path, dest_logo_path)
        print("Fallback border copied to logo path")
    else:
        print("WARNING: Logo not found")

# Copy original diagrams from extracted_img_dir (image3, image4, image5)
orig_diagrams = {
    "img_5_image3.png": "image3.png",
    "img_3_image4.png": "image4.png",
    "img_2_image5.jpeg": "image5.jpeg"
}
for src, dest in orig_diagrams.items():
    src_path = os.path.join(extracted_img_dir, src)
    dest_path = os.path.join(dest_img_dir, dest)
    if os.path.exists(src_path):
        shutil.copy(src_path, dest_path)
        print(f"Copied original diagram {src} to {dest}")
    else:
        print(f"WARNING: Original diagram {src} not found in {src_path}")

# Copy diagrams
for src_name, dest_name in diagram_files.items():
    src_path = os.path.join(scratch_img_dir, src_name)
    dest_path = os.path.join(dest_img_dir, dest_name)
    if os.path.exists(src_path):
        shutil.copy(src_path, dest_path)
        print(f"Copied {src_name} to {dest_name}")
    else:
        print(f"WARNING: Diagram {src_name} not found in {src_path}")

# Initialize Document
doc = docx.Document()

# Helper function to add page number to footer
def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# Function to enforce Times New Roman on a run for all scripts (ASCII, HAnsi, EastAsia, CS)
def enforce_times_new_roman_on_run(run):
    rPr = run._r.get_or_add_rPr()
    rFonts_list = rPr.xpath('./w:rFonts')
    if rFonts_list:
        rFonts = rFonts_list[0]
    else:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

# Function to enforce Times New Roman on a style's default formatting
def enforce_times_new_roman_on_style(style):
    rPr = style._element.get_or_add_rPr()
    rFonts_list = rPr.xpath('./w:rFonts')
    if rFonts_list:
        rFonts = rFonts_list[0]
    else:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')

# Helper to add native page border (double-line) to a section
def add_section_page_border(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'double') # Double border line
        border.set(qn('w:sz'), '12') # Thickness (1.5 pt)
        border.set(qn('w:space'), '24') # Space from edge
        border.set(qn('w:color'), '000000') # Black color
        pgBorders.append(border)
    sectPr.append(pgBorders)

# Helper to clear page borders on a section
def clear_section_page_border(section):
    sectPr = section._sectPr
    pgBorders_list = sectPr.xpath('./w:pgBorders')
    if pgBorders_list:
        pgBorders = pgBorders_list[0]
        pgBorders.clear()
    else:
        pgBorders = OxmlElement('w:pgBorders')
        sectPr.append(pgBorders)
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none') # Explicitly no border
        pgBorders.append(border)

# 2. Page Setup for Section 1 (Cover Page)
cover_section = doc.sections[0]
cover_section.top_margin = Cm(2.0)
cover_section.bottom_margin = Cm(2.0)
cover_section.left_margin = Cm(3.0)
cover_section.right_margin = Cm(2.0)

# Apply native page borders to cover page
add_section_page_border(cover_section)

# Configure Styles
# Base Normal Style
normal_style = doc.styles['Normal']
normal_style.font.name = 'Times New Roman'
normal_style.font.size = Pt(13)
normal_style.paragraph_format.line_spacing = 1.5
normal_style.paragraph_format.space_after = Pt(6)
normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
enforce_times_new_roman_on_style(normal_style)

# Heading 1 Style
h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Times New Roman'
h1_style.font.size = Pt(16)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0, 0, 0)
h1_style.paragraph_format.space_before = Pt(12)
h1_style.paragraph_format.space_after = Pt(6)
h1_style.paragraph_format.keep_with_next = True
h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
enforce_times_new_roman_on_style(h1_style)

# Heading 2 Style
h2_style = doc.styles['Heading 2']
h2_style.font.name = 'Times New Roman'
h2_style.font.size = Pt(14)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0, 0, 0)
h2_style.paragraph_format.space_before = Pt(12)
h2_style.paragraph_format.space_after = Pt(6)
h2_style.paragraph_format.keep_with_next = True
h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
enforce_times_new_roman_on_style(h2_style)

# Heading 3 Style
h3_style = doc.styles['Heading 3']
h3_style.font.name = 'Times New Roman'
h3_style.font.size = Pt(13)
h3_style.font.bold = True
h3_style.font.italic = True
h3_style.font.color.rgb = RGBColor(0, 0, 0)
h3_style.paragraph_format.space_before = Pt(6)
h3_style.paragraph_format.space_after = Pt(6)
h3_style.paragraph_format.keep_with_next = True
h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
enforce_times_new_roman_on_style(h3_style)

# Configure custom styles for Captions (to support automatic TOC generation)
def get_or_create_caption_style(name, is_bold=False, is_italic=True, font_size=11, align=None):
    try:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles['Normal']
    except ValueError:
        style = doc.styles[name]
    
    style.font.name = 'Times New Roman'
    style.font.size = Pt(font_size)
    style.font.bold = is_bold
    style.font.italic = is_italic
    style.font.color.rgb = RGBColor(0, 0, 0)
    if align is not None:
        style.paragraph_format.alignment = align
    enforce_times_new_roman_on_style(style)
    return style

style_fig = get_or_create_caption_style('Caption Hình', is_bold=False, is_italic=True, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
style_tbl = get_or_create_caption_style('Caption Bảng', is_bold=True, is_italic=False, font_size=12)

# Helper function to add TOC
def add_toc(p):
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    
    # Add a run inside fldSimple containing placeholder text
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set font to Times New Roman
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    
    # Set text color to gray to indicate placeholder
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '555555')
    rPr.append(color)
    
    run.append(rPr)
    
    t = OxmlElement('w:t')
    t.text = "Nhấp chuột phải vào dòng chữ này và chọn 'Update Field' để hiển thị mục lục tự động"
    run.append(t)
    
    fldSimple.append(run)
    p._p.append(fldSimple)

# Helper function to add List of Tables / Figures automatically
def add_list_of_captions(p, style_name):
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), f'TOC \\h \\z \\t "{style_name},1"')
    
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '555555')
    rPr.append(color)
    run.append(rPr)
    
    t = OxmlElement('w:t')
    label = "hình ảnh" if "Hình" in style_name else "bảng biểu"
    t.text = f"Nhấp chuột phải vào dòng chữ này và chọn 'Update Field' để hiển thị danh mục {label} tự động"
    run.append(t)
    
    fldSimple.append(run)
    p._p.append(fldSimple)

# Custom styles or helper formatting
def add_custom_para(text="", style_name="Normal", space_after=Pt(6), line_spacing=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False, font_size=13):
    lines = text.split('\n')
    last_p = None
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = line_spacing
        p.paragraph_format.space_after = space_after
        p.paragraph_format.alignment = align
        
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        enforce_times_new_roman_on_run(run)
        last_p = p
    return last_p

def add_heading_1(text):
    p = doc.add_paragraph(text, style='Heading 1')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        enforce_times_new_roman_on_run(run)
    return p

def add_heading_2(text):
    p = doc.add_paragraph(text, style='Heading 2')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        enforce_times_new_roman_on_run(run)
    return p

def add_heading_3(text):
    p = doc.add_paragraph(text, style='Heading 3')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        enforce_times_new_roman_on_run(run)
    return p

def add_figure(img_name, caption):
    img_path = os.path.join(dest_img_dir, img_name)
    if os.path.exists(img_path):
        try:
            with Image.open(img_path) as img:
                w_px, h_px = img.size
            aspect = w_px / h_px
            if w_px >= h_px:
                # Landscape
                width_to_set = Cm(14.5)
            else:
                # Portrait - cap height at 19.5 cm to leave room for caption
                max_height = 19.5
                width_cm = max_height * aspect
                # Cap width at 14.5 cm just in case
                width_to_set = Cm(min(width_cm, 14.5))
        except Exception as e:
            print(f"Error reading image dimensions for {img_name}: {e}")
            width_to_set = Cm(14.5) # fallback

        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(3)
        p_img.paragraph_format.keep_with_next = True
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=width_to_set)
        
        p_cap = doc.add_paragraph(style='Caption Hình')
        p_cap.paragraph_format.space_after = Pt(6)
        p_cap.paragraph_format.keep_with_next = True
        run_cap = p_cap.add_run(caption)
        enforce_times_new_roman_on_run(run_cap)
    else:
        print(f"WARNING: Image not found for figure: {img_path}")
        add_custom_para(f"[HÌNH ẢNH: {caption} - Đường dẫn {img_name} không tồn tại]", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

# Format table cells with background shading and borders
def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def add_data_dictionary_table(title, columns, data):
    p_title = doc.add_paragraph(style='Caption Bảng')
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(3)
    p_title.paragraph_format.keep_with_next = True
    run_t = p_title.add_run(title)
    enforce_times_new_roman_on_run(run_t)
    
    # Create table
    table = doc.add_table(rows=len(data) + 1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Determine column widths dynamically based on headers and column count
    num_cols = len(columns)
    col_widths = []
    first_col = columns[0].lower()
    
    if num_cols == 3:
        # Abbreviations table
        col_widths = [Cm(1.5), Cm(3.5), Cm(10.5)]
    elif num_cols == 4:
        if "nhóm" in first_col or "nhom" in first_col:
            # Function analysis table
            col_widths = [Cm(3.0), Cm(3.5), Cm(6.0), Cm(3.0)]
        else:
            # Test scenarios table
            col_widths = [Cm(3.5), Cm(5.0), Cm(5.0), Cm(2.0)]
    elif num_cols == 5:
        if "api" in first_col:
            # API Test results table
            col_widths = [Cm(4.0), Cm(1.5), Cm(4.0), Cm(4.5), Cm(1.5)]
        else:
            # Data dictionary table
            col_widths = [Cm(3.0), Cm(2.5), Cm(1.5), Cm(1.5), Cm(7.0)]
    else:
        # Default fallback: distribute evenly
        total_width = 15.5
        col_widths = [Cm(total_width / num_cols)] * num_cols

    # Apply widths to every cell in the table
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths):
                cell.width = col_widths[i]
                
    table.allow_autofit = False
    
    # Format header
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
        set_cell_background(hdr_cells[i], "1A365D") # Navy blue
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
        enforce_times_new_roman_on_run(run)
        
    # Fill data
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = str(text)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(2)
            if col_idx in [0, 4]:  # Name, Description -> Left
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:  # Type, Key, Null -> Center
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                run = p.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                enforce_times_new_roman_on_run(run)
                
            # Zebra striping
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F7FAFC") # light gray

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(6)

# Read original document paragraphs to extract relevant text
orig_doc = None
if os.path.exists(original_docx):
    orig_doc = docx.Document(original_docx)
    print("Loaded original docx with", len(orig_doc.paragraphs), "paragraphs.")
else:
    print("WARNING: Original docx not found at", original_docx)

def get_orig_text_range(start_idx, end_idx):
    if not orig_doc:
        return []
    text_list = []
    for idx in range(start_idx, min(end_idx + 1, len(orig_doc.paragraphs))):
        p = orig_doc.paragraphs[idx]
        if p.text.strip():
            txt = p.text.replace("Nhóm sinh viên thực hiện", "Sinh viên thực hiện")
            txt = txt.replace("nhóm thực hiện", "sinh viên thực hiện")
            txt = txt.replace("Nhóm đã quyết định", "Sinh viên đã quyết định")
            txt = txt.replace("chúng tôi", "sinh viên")
            txt = txt.replace("chúng em", "em")
            text_list.append((txt, p.style.name))
    return text_list

def add_extracted_text(start_idx, end_idx):
    paras = get_orig_text_range(start_idx, end_idx)
    for txt, style_name in paras:
        if "Heading" in style_name:
            continue
        if "Caption" in style_name:
            continue
        align = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_custom_para(txt, align=align)

# ==============================================================================
# DOCUMENT GENERATION START
# ==============================================================================

# ----------------- COVER PAGE -----------------
add_custom_para("TRƯỜNG ĐẠI HỌC BÌNH DƯƠNG", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=14, space_after=Pt(2))
add_custom_para("KHOA CÔNG NGHỆ THÔNG TIN, ROBOT VÀ TRÍ TUỆ NHÂN TẠO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=14, space_after=Pt(24))

p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_logo.paragraph_format.space_after = Pt(24)
if os.path.exists(dest_logo_path):
    p_logo.add_run().add_picture(dest_logo_path, width=Cm(4.5))
else:
    p_logo.add_run("[LOGO TRƯỜNG ĐẠI HỌC BÌNH DƯƠNG]")

add_custom_para("TIỂU LUẬN CUỐI KỲ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=16, space_after=Pt(4))
add_custom_para("Học kỳ 3 – Năm học: 2025 – 2026", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, font_size=13, space_after=Pt(8))
add_custom_para("Học phần: CHUYÊN ĐỀ 1", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=18, space_after=Pt(18))

add_custom_para("ĐỀ TÀI:", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=14, space_after=Pt(2))
add_custom_para("XÂY DỰNG VÀ TRIỂN KHAI HỆ THỐNG MẠNG XÃ HỘI\nKẾT HỢP ĐẶT VÉ DU LỊCH \"TRAVELCONNECT\"\nTRÊN NỀN TẢNG CLOUD PLATFORM", 
                align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=16, space_after=Pt(48))

p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_info.paragraph_format.left_indent = Cm(4)
p_info.paragraph_format.space_after = Pt(36)
run_info = p_info.add_run(
    "Giảng viên hướng dẫn:\tThS. Dương Anh Tuấn\n"
    "Sinh viên thực hiện:\tPhan Đình Luyến\n"
    "Mã số sinh viên:\t\t22050036\n"
    "Lớp:\t\t\t22TH01"
)
run_info.font.name = 'Times New Roman'
run_info.font.size = Pt(13)
run_info.bold = True

add_custom_para("Bình Dương, Tháng 7 năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=13)

# Create Section 2 (Rest of document)
section2 = doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
section2.top_margin = Cm(2.0)
section2.bottom_margin = Cm(2.0)
section2.left_margin = Cm(3.0)
section2.right_margin = Cm(2.0)

# Clear page borders on Section 2
clear_section_page_border(section2)

# Unlink footer from cover page section
section2.footer.is_linked_to_previous = False

# Set up page numbering on Section 2's footer
footer2 = section2.footer
footer_p2 = footer2.paragraphs[0]
footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = footer_p2.add_run()
run2.font.name = 'Times New Roman'
run2.font.size = Pt(11)
add_page_number(run2)
enforce_times_new_roman_on_run(run2)

# Set page number format of section 2 to lowercase Roman numerals starting at 1
sectPr2 = section2._sectPr
existing_pgNumType2 = sectPr2.find(qn('w:pgNumType'))
if existing_pgNumType2 is not None:
    sectPr2.remove(existing_pgNumType2)
pgNumType2 = parse_xml(f'<w:pgNumType {nsdecls("w")} w:fmt="lowerRoman" w:start="1"/>')
sectPr2.append(pgNumType2)

# ----------------- ACKNOWLEDGMENT -----------------
add_heading_1("LỜI CẢM ƠN")
add_custom_para(
    "Lời đầu tiên, em xin gửi lời cảm ơn chân thành đến Trường Đại học Bình Dương, Khoa Công nghệ thông tin, "
    "Robot và Trí tuệ nhân tạo đã tạo điều kiện học tập tốt nhất cho em trong suốt thời gian qua. Học phần Chuyên đề 1 "
    "đã giúp em tiếp cận với các công nghệ lập trình Web Full-stack hiện đại và quy trình phát triển, triển khai phần mềm thực tế."
)
add_custom_para(
    "Đặc biệt, em xin bày tỏ lòng biết ơn sâu sắc nhất tới thầy ThS. Dương Anh Tuấn - người giảng viên đã trực tiếp giảng dạy, "
    "hướng dẫn và truyền đạt kiến thức chuyên môn quý báu. Những định hướng và lời khuyên chân thành của thầy đã giúp em định hình ý tưởng "
    "và hoàn thiện hệ thống phần mềm TravelConnect từ kiến trúc 3 tầng cho đến việc triển khai thành công trên môi trường điện toán đám mây."
)
add_custom_para(
    "Mặc dù đã cố gắng hết sức để hoàn thành đồ án và báo cáo một cách chỉn chu nhất, nhưng do giới hạn về mặt thời gian cũng như kiến thức tích lũy, "
    "hệ thống chắc chắn không tránh khỏi những thiếu sót. Em rất mong nhận được sự cảm thông, những ý kiến đóng góp và nhận xét của thầy để em "
    "có cơ hội cải thiện và phát triển ứng dụng tốt hơn trong tương lai."
)
add_custom_para("Em xin chân thành cảm ơn thầy!", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True)

doc.add_page_break()

# ----------------- GRADE SHEET -----------------
add_heading_1("NHẬN XÉT VÀ CHẤM ĐIỂM CỦA GIẢNG VIÊN")
add_custom_para("Nhận xét của giảng viên hướng dẫn:", bold=True)
p_box = doc.add_paragraph()
p_box.paragraph_format.space_after = Pt(12)
run_box = p_box.add_run(
    "............................................................................................................................................................\n"
    "............................................................................................................................................................\n"
    "............................................................................................................................................................\n"
    "............................................................................................................................................................\n"
    "............................................................................................................................................................\n"
    "............................................................................................................................................................\n"
    "............................................................................................................................................................"
)
run_box.font.name = 'Times New Roman'
run_box.font.size = Pt(13)

add_custom_para("Điểm số đánh giá bằng số: .................... (Bằng chữ: ...........................................................)", bold=True, space_after=Pt(36))
add_custom_para("Bình Dương, ngày ....... tháng ....... năm 2026", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True)
add_custom_para("Giảng viên chấm thi\n(Ký và ghi rõ họ tên)", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, space_after=Pt(48))

doc.add_page_break()

# ----------------- TOC -----------------
add_heading_1("MỤC LỤC")
add_custom_para("[Hướng dẫn: Sau khi mở tệp này bằng Microsoft Word, vui lòng nhấn chuột phải vào vùng mục lục bên dưới và chọn 'Update Field' để hiển thị mục lục tự động với số trang chính xác]", italic=True, font_size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
p_toc = doc.add_paragraph()
add_toc(p_toc)

doc.add_page_break()

# ----------------- LIST OF ABBREVIATIONS & TABLES -----------------
add_heading_1("DANH MỤC CÁC KÝ HIỆU, CÁC CHỮ VIẾT TẮT")
abbr_cols = ["STT", "Ký hiệu / Viết tắt", "Ý nghĩa / Tên tiếng Anh đầy đủ"]
abbr_data = [
    ["1", "API", "Application Programming Interface (Giao diện lập trình ứng dụng)"],
    ["2", "AWS", "Amazon Web Services (Nền tảng điện toán đám mây của Amazon)"],
    ["3", "CLO", "Course Learning Outcomes (Chuẩn đầu ra của học phần)"],
    ["4", "CRUD", "Create, Read, Update, Delete (Các thao tác cơ bản với CSDL: Thêm, Đọc, Sửa, Xóa)"],
    ["5", "CSDL", "Cơ sở dữ liệu (Database)"],
    ["6", "DFD", "Data Flow Diagram (Sơ đồ luồng dữ liệu)"],
    ["7", "EC2", "Elastic Compute Cloud (Dịch vụ máy chủ ảo trên đám mây của AWS)"],
    ["8", "ERD", "Entity Relationship Diagram (Sơ đồ mối quan hệ thực thể)"],
    ["9", "HTML", "HyperText Markup Language (Ngôn ngữ đánh dấu siêu văn bản)"],
    ["10", "JWT", "JSON Web Token (Tiêu chuẩn xác thực người dùng bằng token JSON)"],
    ["11", "KDL", "Khu du lịch"],
    ["12", "MVC", "Model - View - Controller (Kiến trúc phần mềm mẫu)"],
    ["13", "OTP", "One-Time Password (Mật khẩu sử dụng một lần)"],
    ["14", "REST", "Representational State Transfer (Kiến trúc truyền tải trạng thái đại diện)"],
    ["15", "UI/UX", "User Interface / User Experience (Giao diện người dùng / Trải nghiệm người dùng)"],
    ["16", "VPC", "Virtual Private Cloud (Mạng riêng ảo trên đám mây)"]
]
add_data_dictionary_table("Bảng danh mục các chữ viết tắt viết trong báo cáo", abbr_cols, abbr_data)

add_heading_1("DANH MỤC CÁC BẢNG BIỂU")
p_tbls = doc.add_paragraph()
add_list_of_captions(p_tbls, "Caption Bảng")

add_heading_1("DANH MỤC CÁC HÌNH ẢNH")
p_figs = doc.add_paragraph()
add_list_of_captions(p_figs, "Caption Hình")

# Create Section 3 (Main chapters)
section3 = doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
section3.top_margin = Cm(2.0)
section3.bottom_margin = Cm(2.0)
section3.left_margin = Cm(3.0)
section3.right_margin = Cm(2.0)

# Clear page borders on Section 3
clear_section_page_border(section3)

# Unlink footer from previous section
section3.footer.is_linked_to_previous = False

# Set up page numbering on Section 3's footer
footer3 = section3.footer
footer_p3 = footer3.paragraphs[0]
footer_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = footer_p3.add_run()
run3.font.name = 'Times New Roman'
run3.font.size = Pt(11)
add_page_number(run3)
enforce_times_new_roman_on_run(run3)

# Set page number format of section 3 to decimal (Arabic) starting at 1
sectPr3 = section3._sectPr
existing_pgNumType3 = sectPr3.find(qn('w:pgNumType'))
if existing_pgNumType3 is not None:
    sectPr3.remove(existing_pgNumType3)
pgNumType3 = parse_xml(f'<w:pgNumType {nsdecls("w")} w:fmt="decimal" w:start="1"/>')
sectPr3.append(pgNumType3)

# ----------------- CHAPTER 1 -----------------
add_heading_1("CHƯƠNG 1. TRÌNH BÀY BÀI TOÁN VÀ CƠ SỞ LÝ THUYẾT (CLO1)")

add_heading_2("1.1. Mô tả bài toán thực tế")

add_heading_3("1.1.1. Lĩnh vực ứng dụng")
add_custom_para(
    "Nền tảng TravelConnect thuộc lĩnh vực công nghệ du lịch (TravelTech) kết hợp mạng xã hội (Social Media). "
    "Đây là một trong những lĩnh vực phát triển nhanh nhất hiện nay, đáp ứng nhu cầu ngày càng cao của thế hệ du khách hiện đại - "
    "những người muốn tự chủ lịch trình, tìm kiếm thông tin chân thực từ cộng đồng và đặt dịch vụ trực tiếp nhanh chóng."
)

add_heading_3("1.1.2. Mục tiêu hệ thống")
add_extracted_text(111, 124)

add_heading_3("1.1.3. Đối tượng sử dụng hệ thống")
add_custom_para(
    "Hệ thống hướng đến việc kết nối 3 đối tượng sử dụng chính trong một hệ sinh thái đồng nhất:\n"
    "1. Khách du lịch (Du khách tự túc): Là đối tượng người dùng cuối, mong muốn tìm kiếm các đánh giá chân thực về điểm đến, "
    "đăng tải trải nghiệm cá nhân dưới dạng bài viết hình ảnh/video, lập kế hoạch chi tiết, tương tác kết bạn, nhắn tin realtime, "
    "đặt mua vé tham quan hoặc gói dịch vụ và thanh toán tiện lợi qua ví điện tử tích hợp, check-in tại KDL bằng QR Code.\n"
    "2. Đối tác khu du lịch (Doanh nghiệp/Nhà cung cấp dịch vụ): Là các cá nhân, doanh nghiệp sở hữu khu nghỉ dưỡng, danh lam thắng cảnh, "
    "điểm vui chơi giải trí. Họ sử dụng TravelConnect để tạo và quản trị hồ sơ doanh nghiệp (được xác thực bởi Admin), đăng bài viết quảng bá dịch vụ, "
    "quản lý danh sách dịch vụ và giá cả, nhận booking từ khách hàng, xử lý xác nhận/hủy đơn đặt vé, quét mã QR check-in của khách và theo dõi doanh thu thông qua biểu đồ thống kê.\n"
    "3. Quản trị viên (Admin): Là nhân sự vận hành hệ thống, chịu trách nhiệm quản lý danh sách người dùng (kích hoạt/khóa tài khoản), phê duyệt "
    "hoặc từ chối các hồ sơ khu du lịch đăng ký mới, cấu hình các thông số nền tảng (phần trăm hoa hồng booking) và giám sát các giao dịch tài chính toàn hệ thống."
)

add_heading_2("1.2. Khảo sát và phân tích yêu cầu")

add_heading_3("1.2.1. Yêu cầu chức năng")
add_custom_para(
    "Dựa trên khảo sát thực tế và phân tích nghiệp vụ, hệ thống TravelConnect được thiết kế đáp ứng đầy đủ các yêu cầu chức năng nghiệp vụ cốt lõi sau:"
)

func_cols = ["Nhóm chức năng", "Chức năng chi tiết", "Mô tả nghiệp vụ", "Vai trò tương ứng"]
func_data = [
    ["Xác thực tài khoản", "Đăng ký tài khoản", "Đăng ký tài khoản mới bằng Email. Hệ thống gửi mã OTP 6 chữ số qua email để kích hoạt.", "Khách, Khu du lịch"],
    ["Xác thực tài khoản", "Đăng nhập & Quên mật khẩu", "Đăng nhập qua JWT token. Hỗ trợ khôi phục mật khẩu thông qua gửi mã xác thực OTP về Email.", "Tất cả vai trò"],
    ["Quản lý hồ sơ", "Cập nhật thông tin", "Thay đổi ảnh đại diện, thông tin cá nhân, cập nhật ví tiền ảo của hệ thống.", "Tất cả vai trò"],
    ["Mạng xã hội du lịch", "Đăng bài chia sẻ", "Đăng bài viết đính kèm nhiều hình ảnh, gắn thẻ Khu du lịch liên quan và chọn danh mục bài viết.", "Khách, Khu du lịch"],
    ["Mạng xã hội du lịch", "Tương tác cộng đồng", "Thích bài viết, bình luận, lưu trữ bài viết yêu thích, đánh giá số sao kèm nhận xét cho KDL.", "Khách du lịch"],
    ["Nghiệp vụ đặt vé", "Xem và đặt dịch vụ", "Khách du lịch xem danh sách dịch vụ của KDL, chọn số lượng vé, tạo hóa đơn đặt vé.", "Khách du lịch"],
    ["Nghiệp vụ đặt vé", "Thanh toán ví ảo", "Thanh toán trực tiếp bằng số dư ví ảo trong tài khoản, hệ thống trừ tiền và tạo mã vé kèm QR code.", "Khách du lịch"],
    ["Nghiệp vụ đặt vé", "Quét mã QR check-in", "Khu du lịch dùng camera/trình quét tích hợp quét mã QR trên vé của khách để kiểm tra tính hợp lệ và check-in.", "Khu du lịch"],
    ["Giao tiếp realtime", "Nhắn tin & Kết bạn", "Kết bạn giữa các tài khoản, gửi lời mời kết bạn, nhắn tin realtime (text, audio, image) qua Socket.IO.", "Khách, Khu du lịch"],
    ["Quản lý đối tác", "Hồ sơ & Dịch vụ KDL", "Doanh nghiệp cập nhật tên KDL, vị trí địa lý, mô tả, tải ảnh bìa và thêm/sửa/xóa các gói dịch vụ bán vé.", "Khu du lịch"],
    ["Quản lý đối tác", "Thống kê doanh thu", "Xem báo cáo thống kê số lượng đặt vé, doanh thu thu về theo các trạng thái qua biểu đồ.", "Khu du lịch"],
    ["Quản trị hệ thống", "Quản lý người dùng", "Admin xem danh sách tài khoản, khóa hoặc mở khóa tài khoản vi phạm chính sách.", "Quản trị viên (Admin)"],
    ["Quản trị hệ thống", "Duyệt hồ sơ đối tác", "Phê duyệt (verified) hoặc từ chối (rejected) hồ sơ KDL để đảm bảo chất lượng dịch vụ trên sàn.", "Quản trị viên (Admin)"]
]
add_data_dictionary_table("Bảng 1.1: Bảng phân tích yêu cầu chức năng chi tiết", func_cols, func_data)

add_heading_3("1.2.2. Yêu cầu phi chức năng")
add_custom_para(
    "Bên cạnh các tính năng nghiệp vụ, TravelConnect cần đáp ứng các tiêu chuẩn chất lượng kỹ thuật (Non-functional requirements) sau:\n"
    "1. Tính bảo mật (Security): Mật khẩu người dùng bắt buộc phải được băm (hash) bằng thư viện bcrypt trước khi lưu trữ vào MySQL. "
    "Các API được bảo vệ nghiêm ngặt bằng cơ chế JWT, ngăn chặn các truy cập trái phép. Toàn bộ thông tin nhạy cảm của người dùng không được phản hồi dưới dạng text thường.\n"
    "2. Hiệu năng & Tốc độ phản hồi (Performance): Thời gian phản hồi trung bình cho các API RESTful phải dưới 500ms dưới điều kiện tải thông thường. "
    "Giao diện Frontend được tối ưu hóa quá trình render bằng React Virtual DOM, đảm bảo mượt mà và không gián đoạn.\n"
    "3. Khả năng mở rộng (Scalability): Hệ thống được thiết kế theo các container Docker độc lập, dễ dàng nâng cấp (scale-up) tài nguyên phần cứng "
    "hoặc scale-out bằng cách khởi chạy nhiều instance backend đằng sau một Nginx Load Balancer khi lượng truy cập tăng đột biến.\n"
    "4. Tính khả dụng & Khả năng tương thích (Usability & Compatibility): Giao diện ứng dụng phải thiết kế đáp ứng (Responsive Design) hoạt động tốt trên cả trình duyệt máy tính (Desktop), "
    "máy tính bảng (Tablet) và điện thoại di động (Smartphones), hỗ trợ hiển thị tốt trên các trình duyệt phổ biến như Chrome, Safari, Firefox, Edge."
)

add_heading_2("1.3. Cơ sở lý thuyết")

add_heading_3("1.3.1. Kiến trúc 3 tầng (3-tier Architecture)")
add_custom_para(
    "Kiến trúc 3 tầng là mô hình phân rã phần mềm kinh duyệt, giúp phân tách rõ ràng nhiệm vụ của từng lớp công nghệ, từ đó nâng cao tính độc lập, "
    "dễ bảo trì và bảo mật của ứng dụng Web. Hệ thống TravelConnect tuân thủ nghiêm ngặt mô hình này:"
)
add_custom_para(
    "- Lớp trình diễn (Presentation Layer / Frontend): Là giao diện người dùng hiển thị trên trình duyệt. Lớp này được phát triển bằng React và Vite, "
    "chịu trách nhiệm hiển thị cấu trúc dữ liệu, nhận thao tác chuột, bàn phím của người dùng, thực hiện gọi API (giao tiếp HTTP) và render lại giao diện động."
)
add_custom_para(
    "- Lớp logic nghiệp vụ (Application Layer / Backend API): Trọng tâm xử lý của ứng dụng. Được phát triển bằng Node.js và Express.js, lớp này chịu trách nhiệm "
    "tiếp nhận yêu cầu từ Presentation Layer, kiểm tra tính hợp lệ của dữ liệu, thực hiện tính toán nghiệp vụ (xác thực mật khẩu, trừ tiền ví, tạo mã QR) và "
    "giao tiếp với Database Layer."
)
add_custom_para(
    "- Lớp dữ liệu (Data Layer / Database Server): Nơi lưu trữ vĩnh viễn dữ liệu hệ thống. Sử dụng MySQL 8, chịu trách nhiệm lưu trữ các bảng cơ sở dữ liệu quan hệ, "
    "quản lý các ràng buộc khóa ngoại, thực thi các truy vấn SQL để thêm, sửa, xóa hoặc lấy dữ liệu theo yêu cầu từ Backend API."
)
add_figure("image3.png", "Hình 1.1: Sơ đồ mô hình kiến trúc Client - Server của hệ thống Web")

add_heading_3("1.3.2. Kiến trúc RESTful API")
add_custom_para(
    "RESTful API là một phong cách kiến trúc phần mềm dựa trên các tiêu chuẩn web công cộng, sử dụng HTTP làm giao thức truyền thông chính. "
    "REST định nghĩa các tài nguyên (resources) thông qua các URI cụ thể và thao tác trên chúng bằng các phương thức HTTP tiêu chuẩn:\n"
    "- GET: Lấy thông tin tài nguyên (không làm thay đổi trạng thái hệ thống).\n"
    "- POST: Tạo mới một tài nguyên.\n"
    "- PUT / PATCH: Cập nhật tài nguyên đã tồn tại.\n"
    "- DELETE: Xóa tài nguyên khỏi hệ thống.\n"
    "TravelConnect triển khai RESTful API cho phép Frontend giao tiếp bất đồng bộ (Asynchronous) với Backend mà không cần tải lại trang, "
    "truyền tải dữ liệu qua định dạng JSON nhẹ nhàng và tối ưu băng thông."
)
add_figure("image5.jpeg", "Hình 1.2: Mô hình kết nối RESTful API và tích hợp dịch vụ bên thứ ba")

add_heading_3("1.3.3. Cơ chế Xác thực và Phân quyền (Authentication / Authorization)")
add_custom_para(
    "Xác thực (Authentication) là quá trình kiểm tra danh tính của người dùng (họ có đúng là người họ khai báo hay không). "
    "Hệ thống sử dụng cơ chế xác thực Token-based với JWT (JSON Web Token). Khi người dùng đăng nhập thành công, Backend sinh ra một chuỗi mã hóa ký số JWT chứa ID tài khoản và vai trò của họ. "
    "Chuỗi này được Frontend lưu trữ và tự động đính kèm vào header `Authorization: Bearer <token>` trong mọi request tiếp theo."
)
add_custom_para(
    "Phân quyền (Authorization) là quá trình xác định tài khoản đã xác thực có quyền thực hiện hành động cụ thể hay không. "
    "TravelConnect phân quyền rõ ràng qua 3 vai trò: `khach_du_lich`, `khu_du_lich` và `admin`. Các middleware trên Backend sẽ giải mã token JWT, "
    "kiểm tra trường `vai_tro` trước khi cho phép request tiếp cận vào Controller xử lý logic nghiệp vụ nhạy cảm."
)

add_heading_3("1.3.4. Thiết kế cơ sở dữ liệu (Database Design)")
add_custom_para(
    "Thiết kế cơ sở dữ liệu quan hệ đóng vai trò quyết định đến tính toàn vẹn của dữ liệu và hiệu năng truy vấn. Hệ thống sử dụng MySQL 8, "
    "áp dụng các nguyên tắc chuẩn hóa (Normalized Forms - 1NF, 2NF, 3NF) nhằm loại bỏ sự dư thừa dữ liệu và phòng ngừa các lỗi bất thường khi thêm/sửa/xóa (Anomaly). "
    "Các bảng được liên kết chặt chẽ thông qua các ràng buộc khóa ngoại (Foreign Keys) với tùy chọn `ON DELETE CASCADE` hoặc `ON DELETE SET NULL` "
    "đảm bảo tính nhất quán tham chiếu dữ liệu tự động."
)

add_heading_3("1.3.5. Điện toán đám mây (Cloud Computing) và Containerization")
add_custom_para(
    "Điện toán đám mây cung cấp các dịch vụ tài nguyên công nghệ thông tin theo nhu cầu qua Internet với mô hình thanh toán theo mức sử dụng. "
    "Ứng dụng TravelConnect tận dụng dịch vụ IaaS (Infrastructure as a Service) của Amazon Web Services (AWS EC2) để chạy máy chủ ảo Ubuntu Server."
)
add_custom_para(
    "Để đơn giản hóa quy trình đóng gói và triển khai ứng dụng nhất quán giữa môi trường máy cục bộ (development) và máy chủ đám mây (production), "
    "ứng dụng sử dụng công nghệ Docker để đóng gói các thành phần Frontend, Backend và Database vào các Container cô lập. "
    "Docker Compose điều phối các container này, thiết lập mạng nội bộ (network) và phân vùng lưu trữ bền vững (volume) giúp hệ thống vận hành trơn tru "
    "chỉ với một lệnh khởi chạy duy nhất."
)
add_figure("image4.png", "Hình 1.3: Sơ đồ kiến trúc Microservices và Container hóa trên Cloud")

doc.add_page_break()

# ----------------- CHAPTER 2 -----------------
add_heading_1("CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG (CLO2)")

add_heading_2("2.1. Phân tích hệ thống")

add_heading_3("2.1.1. Sơ đồ Use Case tổng quát")
add_custom_para(
    "Sơ đồ Use Case thể hiện trực quan các chức năng của hệ thống dưới góc nhìn tương tác giữa các tác nhân (Actors) bao gồm Khách du lịch, Đối tác khu du lịch và Quản trị viên (Admin) với các nghiệp vụ."
)
add_figure("usecase.png", "Hình 2.1: Sơ đồ Use Case tổng quát hệ thống TravelConnect")

add_heading_3("2.1.2. Đặc tả chức năng hệ thống")
add_custom_para(
    "Để làm rõ luồng nghiệp vụ trong sơ đồ Use Case, dưới đây là đặc tả chi tiết của một số chức năng cốt lõi:"
)
add_custom_para(
    "1. Đặc tả chức năng Đăng ký tài khoản:\n"
    "- Tác nhân kích hoạt: Người dùng chưa có tài khoản (Khách du lịch hoặc Đối tác KDL).\n"
    "- Tiền điều kiện: Địa chỉ Email đăng ký chưa tồn tại trên hệ thống.\n"
    "- Luồng xử lý chính:\n"
    "  1. Người dùng nhập thông tin đăng ký (Họ tên, Email, Mật khẩu, Vai trò mong muốn) và nhấn gửi.\n"
    "  2. Hệ thống kiểm tra trùng lặp email. Nếu không trùng, tạo tài khoản trạng thái chưa xác thực (`da_xac_thuc_otp = 0`).\n"
    "  3. Hệ thống tạo mã OTP 6 chữ số ngẫu nhiên, lưu vào bảng `otp_xac_thuc` kèm thời gian hết hạn (5 phút) và gửi OTP qua email cho người dùng.\n"
    "  4. Người dùng nhận email, nhập mã OTP vào giao diện kích hoạt trên Web.\n"
    "  5. Hệ thống đối chiếu OTP. Nếu hợp lệ, cập nhật `da_xac_thuc_otp = 1` và chuyển hướng sang màn hình đăng nhập.\n"
    "- Hậu điều kiện: Tài khoản được kích hoạt thành công và có thể đăng nhập."
)
add_custom_para(
    "2. Đặc tả chức năng Đặt vé và Thanh toán ví:\n"
    "- Tác nhân kích hoạt: Khách du lịch đã đăng nhập.\n"
    "- Tiền điều kiện: Số dư ví ảo của du khách phải lớn hơn hoặc bằng tổng tiền thanh toán của hóa đơn.\n"
    "- Luồng xử lý chính:\n"
    "  1. Du khách duyệt hồ sơ khu du lịch, xem danh sách gói dịch vụ và chọn đặt chỗ.\n"
    "  2. Nhập số lượng khách, ngày đến và nhấn 'Thanh toán'.\n"
    "  3. Hệ thống tạo hóa đơn trạng thái pending. Thực hiện kiểm tra số dư ví ảo.\n"
    "  4. Nếu số dư hợp lệ, hệ thống thực hiện trừ tiền ví khách hàng, cộng tiền ví ảo cho KDL (sau khi trừ chiết khấu hoa hồng của hệ thống).\n"
    "  5. Cập nhật trạng thái hóa đơn thành completed, tạo đơn đặt vé trạng thái pending và sinh mã QR Code chứa thông tin ID đặt vé để check-in.\n"
    "- Hậu điều kiện: Số dư ví thay đổi, vé được tạo kèm QR Code, thông báo tự động được gửi tới KDL."
)

add_heading_2("2.2. Thiết kế kiến trúc hệ thống")

add_heading_3("2.2.1. Kiến trúc tổng thể")
add_custom_para(
    "Hệ thống TravelConnect được xây dựng theo mô hình kiến trúc 3 tầng chuẩn, kết hợp Web client SPA (Single Page Application) "
    "giao tiếp với Backend API qua giao thức HTTP RESTful và giao tiếp Realtime bằng Socket.IO (WebSockets)."
)
add_figure("dfd_context.png", "Hình 2.2: Sơ đồ luồng dữ liệu (DFD) mức ngữ cảnh")
add_figure("dfd_level1.png", "Hình 2.3: Sơ đồ luồng dữ liệu (DFD) mức 1")

add_heading_3("2.2.2. Kiến trúc Frontend (Frontend Architecture)")
add_custom_para(
    "Frontend của TravelConnect được phát triển trên thư viện React kết hợp công cụ build siêu tốc Vite. Cấu trúc thư mục được thiết kế khoa học:"
)
add_custom_para(
    "- `/src/components`: Chứa các thành phần giao diện dùng chung (Button, Input, Sidebar, Modal, Navbar).\n"
    "- `/src/pages`: Chứa giao diện các trang độc lập (Home, Login, Register, Explore, Booking, Profile, AdminDashboard, BusinessDashboard).\n"
    "- `/src/context`: Chứa các React Context quản lý trạng thái toàn cục như thông tin đăng nhập (`AuthContext`), giỏ hàng, thông báo.\n"
    "- `/src/services`: Định nghĩa các hàm gọi API sử dụng Axios, đính kèm Token JWT tự động vào request interceptor.\n"
    "- `/src/assets`: Chứa hình ảnh, CSS và tài nguyên tĩnh của ứng dụng."
)

add_heading_3("2.2.3. Kiến trúc Backend (Backend Architecture)")
add_custom_para(
    "Backend được xây dựng bằng Node.js và Express.js theo kiến trúc MVC định hướng API (Controller-Service-Repository). Cấu trúc thư mục gồm:"
)
add_custom_para(
    "- `/config`: Cấu hình kết nối cơ sở dữ liệu MySQL (`db.js`), cấu hình transporter gửi mail (`nodemailer.js`).\n"
    "- `/controllers`: Xử lý trực tiếp các yêu cầu HTTP, nhận dữ liệu đầu vào, gọi Service và trả về kết quả JSON.\n"
    "- `/middlewares`: Chứa các bộ lọc request (ví dụ: `authMiddleware.js` giải mã kiểm tra token JWT, `roleMiddleware.js` phân quyền, upload file bằng multer).\n"
    "- `/routes`: Định nghĩa định tuyến API phân nhóm như `/api/auth`, `/api/posts`, `/api/bookings`, `/api/payments`, v.v.\n"
    "- `/uploads`: Thư mục lưu trữ các file đa phương tiện (ảnh bài viết, avatar) tải lên từ Client.\n"
    "- `server.js`: Điểm khởi chạy ứng dụng, thiết lập Express server, tích hợp Socket.IO và lắng nghe cổng kết nối."
)

add_heading_3("2.2.4. Kiến trúc Database (Database Architecture)")
add_custom_para(
    "Hệ quản trị CSDL quan hệ MySQL 8.0 đảm nhận vai trò quản trị lưu trữ. Database bao gồm 9 bảng được thiết kế liên kết chặt chẽ bằng khóa ngoại, "
    "được đánh chỉ mục (Index) trên các cột thường dùng để tìm kiếm như `email`, `trang_thai`, `ngay_den` nhằm tối ưu tốc độ truy vấn SQL."
)

add_heading_2("2.3. Thiết kế cơ sở dữ liệu")

add_heading_3("2.3.1. Sơ đồ quan hệ thực thể (ERD)")
add_custom_para(
    "Sơ đồ ERD (Entity Relationship Diagram) dưới đây biểu diễn cấu trúc các thực thể và các mối quan hệ (1-1, 1-n, n-n) trong cơ sở dữ liệu TravelConnect."
)
add_figure("erd.png", "Hình 2.4: Sơ đồ thực thể liên kết ERD hệ thống TravelConnect")

add_heading_3("2.3.2. Từ điển dữ liệu (Data Dictionary)")
add_custom_para(
    "Dưới đây là đặc tả chi tiết của 5 bảng cơ sở dữ liệu quan trọng nhất trong hệ thống TravelConnect:"
)

# 1. Bảng nguoi_dung
nd_cols = ["Tên cột (Field)", "Kiểu dữ liệu", "Khóa (Key)", "Cho phép Null", "Mô tả ý nghĩa"]
nd_data = [
    ["id", "INT AUTO_INCREMENT", "PRIMARY KEY", "NO", "Mã định danh duy nhất của người dùng"],
    ["ten", "VARCHAR(100)", "-", "NO", "Họ tên người dùng hiển thị"],
    ["email", "VARCHAR(100)", "UNIQUE", "NO", "Email dùng để đăng nhập và nhận OTP"],
    ["mat_khau", "VARCHAR(255)", "-", "NO", "Mật khẩu đã được mã hóa bcrypt"],
    ["anh_dai_dien", "VARCHAR(255)", "-", "YES", "Đường dẫn file ảnh avatar của người dùng"],
    ["vai_tro", "ENUM('khach_du_lich','khu_du_lich','admin')", "-", "NO", "Phân vai trò tài khoản trong hệ thống"],
    ["diem_tin_cay", "INT", "-", "YES", "Điểm uy tín tích lũy (mặc định 50)"],
    ["da_xac_thuc_otp", "TINYINT(1)", "-", "YES", "Đánh dấu tài khoản đã xác thực OTP (0/1)"],
    ["so_du", "DECIMAL(15,0)", "-", "YES", "Số dư ví ảo dùng thanh toán vé"],
    ["trang_thai_tai_khoan", "ENUM('active','suspended')", "-", "YES", "Trạng thái hoạt động của tài khoản"],
    ["so_thich_json", "JSON", "-", "YES", "Mảng chứa danh mục sở thích du lịch"]
]
add_data_dictionary_table("Bảng 2.1: Từ điển dữ liệu bảng nguoi_dung (người dùng)", nd_cols, nd_data)

# 2. Bảng ho_so_khu_du_lich
kdl_cols = ["Tên cột (Field)", "Kiểu dữ liệu", "Khóa (Key)", "Cho phép Null", "Mô tả ý nghĩa"]
kdl_data = [
    ["id_nguoi_dung", "INT", "PRIMARY, FOREIGN KEY", "NO", "Liên kết với id trong bảng nguoi_dung"],
    ["ten_khu_du_lich", "VARCHAR(255)", "-", "NO", "Tên khu du lịch hoặc tên doanh nghiệp"],
    ["giay_phep_kinh_doanh", "VARCHAR(255)", "-", "YES", "Mã số thuế hoặc giấy phép kinh doanh"],
    ["dia_chi_chi_tiet", "TEXT", "-", "YES", "Địa chỉ cụ thể của khu du lịch"],
    ["tinh_thanh", "VARCHAR(100)", "-", "YES", "Tỉnh/Thành phố tọa lạc"],
    ["mo_ta_tong_quan", "TEXT", "-", "YES", "Bài viết mô tả giới thiệu về KDL"],
    ["vi_do", "FLOAT", "-", "YES", "Tọa độ vĩ độ dùng hiển thị bản đồ"],
    ["kinh_do", "FLOAT", "-", "YES", "Tọa độ kinh độ dùng hiển thị bản đồ"],
    ["trang_thai_duyet", "ENUM('pending','verified','rejected')", "-", "YES", "Trạng thái phê duyệt của Admin"],
    ["hinh_anh_bia", "VARCHAR(255)", "-", "YES", "Ảnh bìa của khu du lịch"]
]
add_data_dictionary_table("Bảng 2.2: Từ điển dữ liệu bảng ho_so_khu_du_lich (hồ sơ khu du lịch)", kdl_cols, kdl_data)

# 3. Bảng bai_viet
bv_cols = ["Tên cột (Field)", "Kiểu dữ liệu", "Khóa (Key)", "Cho phép Null", "Mô tả ý nghĩa"]
bv_data = [
    ["id", "INT AUTO_INCREMENT", "PRIMARY KEY", "NO", "Mã định danh duy nhất của bài viết"],
    ["id_nguoi_dung", "INT", "FOREIGN KEY", "YES", "Liên kết id người đăng bài"],
    ["tieu_de", "VARCHAR(255)", "-", "YES", "Tiêu đề bài viết"],
    ["noi_dung", "TEXT", "-", "YES", "Nội dung văn bản chia sẻ"],
    ["hinh_anh_json", "JSON", "-", "YES", "Mảng chứa danh sách đường dẫn ảnh đính kèm"],
    ["loai_bai_viet", "ENUM('chia_se','quang_cao')", "-", "YES", "Phân loại bài viết chia sẻ hay quảng cáo"],
    ["luot_thich", "INT", "-", "YES", "Số lượng lượt thích bài viết (mặc định 0)"],
    ["id_kdl_gan_the", "INT", "FOREIGN KEY", "YES", "Liên kết gắn thẻ khu du lịch"],
    ["danh_muc", "VARCHAR(100)", "-", "YES", "Danh mục phân loại (Ẩm thực, Khám phá, Nghỉ dưỡng)"]
]
add_data_dictionary_table("Bảng 2.3: Từ điển dữ liệu bảng bai_viet (bài viết chia sẻ)", bv_cols, bv_data)

# 4. Bảng dat_ve
dv_cols = ["Tên cột (Field)", "Kiểu dữ liệu", "Khóa (Key)", "Cho phép Null", "Mô tả ý nghĩa"]
dv_data = [
    ["id", "INT AUTO_INCREMENT", "PRIMARY KEY", "NO", "Mã đặt vé duy nhất trong hệ thống"],
    ["id_kdl", "INT", "FOREIGN KEY", "NO", "ID của khu du lịch cung cấp vé"],
    ["id_khach", "INT", "FOREIGN KEY", "NO", "ID của khách du lịch đặt vé"],
    ["ngay_den", "DATETIME", "-", "NO", "Ngày giờ khách dự kiến đến check-in"],
    ["so_ngay", "INT", "-", "YES", "Số lượng ngày lưu trú (mặc định 1)"],
    ["so_nguoi", "INT", "-", "YES", "Số lượng khách tham gia"],
    ["loai_ve", "VARCHAR(50)", "-", "YES", "Loại hình dịch vụ vé chọn"],
    ["tong_tien", "DECIMAL(15,2)", "-", "YES", "Tổng số tiền thanh toán"],
    ["trang_thai", "ENUM('pending','confirmed','completed','cancelled')", "-", "YES", "Trạng thái đặt vé"],
    ["ghi_chu", "TEXT", "-", "YES", "Ghi chú thêm của khách du lịch"],
    ["ngay_tao", "TIMESTAMP", "-", "YES", "Thời gian khởi tạo đơn đặt vé"]
]
add_data_dictionary_table("Bảng 2.4: Từ điển dữ liệu bảng dat_ve (đơn đặt vé)", dv_cols, dv_data)

# 5. Bảng thanh_toan
tt_cols = ["Tên cột (Field)", "Kiểu dữ liệu", "Khóa (Key)", "Cho phép Null", "Mô tả ý nghĩa"]
tt_data = [
    ["id", "INT AUTO_INCREMENT", "PRIMARY KEY", "NO", "Mã hóa đơn thanh toán duy nhất"],
    ["ma_tra_cuu", "VARCHAR(20)", "UNIQUE", "YES", "Mã tra cứu hóa đơn dạng chuỗi"],
    ["id_nguoi_dung", "INT", "FOREIGN KEY", "NO", "ID người thực hiện thanh toán"],
    ["id_kdl", "INT", "FOREIGN KEY", "YES", "ID khu du lịch thụ hưởng"],
    ["tong_tien", "DECIMAL(15,0)", "-", "NO", "Số tiền thanh toán hóa đơn"],
    ["trang_thai", "ENUM('pending','completed','cancelled')", "-", "YES", "Trạng thái hóa đơn"],
    ["phuong_thuc", "VARCHAR(50)", "-", "YES", "Phương thức thanh toán (ví dụ: wallet)"],
    ["ngay_tao", "TIMESTAMP", "-", "YES", "Thời gian giao dịch diễn ra"]
]
add_data_dictionary_table("Bảng 2.5: Từ điển dữ liệu bảng thanh_toan (giao dịch thanh toán)", tt_cols, tt_data)

add_heading_2("2.4. Thiết kế giao diện và luồng xử lý chính")

add_heading_3("2.4.1. Thiết kế Giao diện (Mockup / Wireframe)")
add_custom_para(
    "Giao diện của TravelConnect được thiết kế hiện đại, sử dụng phong cách Glassmorphism và tối ưu hóa trải nghiệm tương tác trực quan:\n"
    "1. Trang khám phá (Feed): Thiết kế theo dạng dòng thời gian (timeline). Bài viết được hiển thị với khung ảnh lớn, "
    "nút thích, bình luận, chia sẻ và đánh giá số sao được bố trí trực quan. Có thanh bên trái điều hướng nhanh và thanh bên phải "
    "gợi ý các khu du lịch nổi bật dựa trên điểm đánh giá cao.\n"
    "2. Trang chi tiết Khu du lịch và Đặt vé: Hiển thị ảnh bìa rộng (hero banner), mô tả thông tin, vị trí trên bản đồ tương tác, "
    "danh sách gói dịch vụ kèm giá bán rõ ràng. Du khách có thể chọn nhanh số lượng vé, chọn ngày đi bằng date-picker và thực hiện đặt vé qua hộp thoại (modal) thanh toán chỉ trong 1 click.\n"
    "3. Dashboard Đối tác và Admin: Thiết kế dạng lưới (Grid) chuyên nghiệp với thanh sidebar cố định. Hiển thị các khối chỉ số tổng quan (số người dùng, số bài viết, doanh thu ví, tổng đơn đặt) "
    "dưới dạng biểu đồ cột và biểu đồ đường trực quan, giúp đối tác và admin dễ dàng theo dõi chỉ số sức khỏe của hệ thống."
)

add_heading_3("2.4.2. Luồng xử lý chính của người dùng (User Flow)")
add_custom_para(
    "Sơ đồ hành trình luồng đi của người dùng (User Flow) mô tả các bước tương tác từ khi truy cập hệ thống đến khi hoàn tất một chu trình đặt dịch vụ du lịch."
)
add_figure("userflow.png", "Hình 2.5: Sơ đồ luồng đi của người dùng (User Flow)")

add_heading_3("2.4.3. Các sơ đồ luồng hoạt động chi tiết")
add_custom_para(
    "Để đảm bảo lập trình chính xác, dưới đây là các sơ đồ luồng hoạt động (Activity Diagrams) mô tả thuật toán xử lý của hệ thống:"
)
add_figure("flow_register.png", "Hình 2.6: Sơ đồ luồng đăng ký tài khoản và xác thực mã OTP email")
add_figure("flow_booking.png", "Hình 2.7: Sơ đồ luồng đặt vé, thanh toán ví ảo và check-in vé bằng QR Code")
add_figure("flow_explore.png", "Hình 2.8: Sơ đồ luồng khám phá bài viết và tương tác của khách du lịch")
add_figure("flow_business.png", "Hình 2.9: Sơ đồ luồng quản lý hồ sơ và dịch vụ của đối tác khu du lịch")

doc.add_page_break()

# ----------------- CHAPTER 3 -----------------
add_heading_1("CHƯƠNG 3. XÂY DỰNG VÀ ĐÁNH GIÁ HỆ THỐNG (CLO2)")

add_heading_2("3.1. Xây dựng Frontend")
add_custom_para(
    "Frontend của ứng dụng TravelConnect được hiện thực hoàn chỉnh dựa trên thư viện ReactJS, sử dụng Vite làm môi trường phát triển nhanh. "
    "Giao diện được lập trình với các đặc điểm nổi bật:\n"
    "1. Giao diện SPA mượt mà: Sử dụng `react-router-dom` để định tuyến các trang client-side, giúp việc chuyển trang diễn ra lập tức, "
    "mang lại cảm giác mượt mà của một ứng dụng native.\n"
    "2. Giao diện Responsive Design: Toàn bộ CSS được viết bằng Grid và Flexbox kết hợp Media Queries chuẩn, "
    "đảm bảo giao diện co giãn hoàn hảo từ màn hình máy tính rộng (1920px) xuống màn hình điện thoại thông minh nhỏ nhất (320px). "
    "Thanh menu điều hướng trên mobile được thu gọn thành biểu tượng menu rút gọn (hamburger menu) tiện lợi.\n"
    "3. Quét QR bằng Camera: Tích hợp thư viện quét mã QR trên Frontend, cho phép đối tác khu du lịch mở trực tiếp camera điện thoại "
    "hoặc laptop để quét mã QR check-in vé của khách ngay tại quầy kiểm soát vé."
)
add_custom_para(
    "Dưới đây là hình ảnh giao diện thực tế của Trang chủ TravelConnect hiển thị dòng tin bài viết chia sẻ trải nghiệm du lịch "
    "và khám phá các khu du lịch nổi bật:"
)
add_figure("screenshot_homepage.png", "Hình 3.1: Giao diện Trang chủ của Nền tảng du lịch kết hợp mạng xã hội TravelConnect")

add_heading_2("3.2. Xây dựng Backend")
add_custom_para(
    "Backend API được lập trình bằng Node.js và Express.js, đảm nhận toàn bộ business logic và các API endpoints của hệ thống:\n"
    "1. Nhóm API chính: Cung cấp đầy đủ các endpoint theo chuẩn RESTful bao gồm xác thực `/api/auth`, bài viết `/api/posts`, "
    "đặt vé `/api/bookings`, thanh toán ví `/api/payments`, chat realtime `/api/messages`, hồ sơ doanh nghiệp `/api/businesses` và quản trị `/api/admin`.\n"
    "2. Xử lý ví điện tử hệ thống: Logic trừ tiền ví ảo của khách du lịch và cộng tiền cho khu du lịch thụ hưởng được bao bọc trong một "
    "Database Transaction (giao dịch CSDL) nhằm đảm bảo tính toàn vẹn dữ liệu. Nếu bất kỳ lỗi nào xảy ra trong quá trình cập nhật số dư, "
    "toàn bộ giao dịch sẽ được Rollback về trạng thái cũ để tránh thất thoát tài chính.\n"
    "3. Tích hợp Realtime bằng Socket.IO: Thiết lập kết nối hai chiều (Bi-directional) thời gian thực. Khi người dùng nhắn tin, "
    "hệ thống sẽ phát (broadcast) tin nhắn lập tức đến phòng chat tương ứng. Khi có đơn đặt vé mới, hệ thống tự động đẩy thông báo "
    "realtime đến tài khoản đối tác KDL mà không cần họ phải tải lại trang dashboard."
)

add_heading_2("3.3. Xây dựng Cơ sở dữ liệu")
add_custom_para(
    "Cơ sở dữ liệu được khởi tạo hoàn chỉnh thông qua tệp SQL Script `database/travelconnect.sql`. Tệp chứa toàn bộ cấu trúc bảng, "
    "các ràng buộc khóa ngoại (foreign key constraints), các chỉ mục (index) và tập dữ liệu mẫu phong phú."
)
add_custom_para(
    "Dữ liệu mẫu (Seed data) bao gồm:\n"
    "- 5 tài khoản mẫu phân bổ đầy đủ các vai trò: 1 tài khoản Admin (`admin@travelconnect.vn`), 2 tài khoản đối tác KDL "
    "(`Phan Dinh Luyen` sở hữu KDL meo emo, và `Đà Lạt Wonder Resort` sở hữu khu nghỉ dưỡng bên hồ Tuyền Lâm), 2 tài khoản khách du lịch.\n"
    "- Hơn 10 bài viết mẫu chia sẻ trải nghiệm du lịch thực tế tại Đà Lạt kèm danh sách ảnh mẫu.\n"
    "- Các bản ghi hóa đơn thanh toán, lịch sử đặt vé và tin nhắn chat mẫu để đảm bảo hệ thống có dữ liệu trực quan ngay khi triển khai."
)

add_heading_2("3.4. Hiện thực Authentication và Authorization")
add_custom_para(
    "Cơ chế bảo mật của TravelConnect được lập trình chặt chẽ thông qua các bước công nghệ:\n"
    "1. Đăng ký tài khoản và OTP Email: Người dùng gửi yêu cầu đăng ký -> Backend tạo mã OTP 6 số ngẫu nhiên -> Lưu mã hóa OTP kèm thời gian hết hạn -> "
    "Sử dụng thư viện `nodemailer` kết nối qua Gmail SMTP gửi email thật chứa mã kích hoạt về hòm thư người dùng. "
    "Người dùng nhập đúng mã OTP trên Web, tài khoản mới được kích hoạt hoạt động.\n"
    "2. Đăng nhập JWT: Khi người dùng gửi email và mật khẩu -> Backend dùng bcrypt so khớp mật khẩu băm -> Sinh chuỗi token JWT ký số bằng thuật toán HMAC-SHA256 với "
    "khóa bí mật `JWT_SECRET`. Token này được Frontend đính kèm vào header `Authorization: Bearer <token>` ở mỗi request sau đó.\n"
    "3. Phân quyền Middleware: Viết middleware `auth.js` để giải mã token. Viết middleware `checkRole.js` so sánh vai trò người dùng. "
    "Ví dụ, endpoint `/api/admin/*` chỉ cho phép tài khoản có vai trò `admin` truy cập. Nếu tài khoản `khach_du_lich` cố tình truy cập, "
    "hệ thống lập tức trả về mã trạng thái HTTP `403 Forbidden`."
)

add_heading_2("3.5. Kiểm thử hệ thống")

add_heading_3("3.5.1. Kiểm thử API (API Testing)")
add_custom_para(
    "Kiểm thử API được tiến hành bằng công cụ Postman để kiểm tra tính đúng đắn của các đầu endpoint API. Kết quả như sau:"
)

test_api_cols = ["API Endpoint", "Phương thức", "Tham số / Request Body", "Kết quả mong đợi", "Mã trạng thái HTTP"]
test_api_data = [
    ["/api/auth/register", "POST", "email, ten, mat_khau, vai_tro", "Đăng ký thành công, gửi OTP về mail", "201 Created"],
    ["/api/auth/verify-otp", "POST", "email, ma_otp", "Xác thực OTP thành công, kích hoạt nick", "200 OK"],
    ["/api/auth/login", "POST", "email, mat_khau", "Đăng nhập thành công, trả về JWT Token", "200 OK"],
    ["/api/bookings/create", "POST", "id_kdl, ngay_den, loai_ve, tong_tien", "Tạo đơn đặt vé mới, trừ tiền ví ảo", "201 Created"],
    ["/api/admin/users", "GET", "Headers: Bearer <token>", "Trả về danh sách tất cả người dùng hệ thống", "200 OK"],
    ["/api/admin/users", "GET", "Không gửi token / Gửi token khách", "Báo lỗi không có quyền truy cập", "403 Forbidden"]
]
add_data_dictionary_table("Bảng 3.1: Kết quả kiểm thử các API Endpoint cốt lõi", test_api_cols, test_api_data)

add_heading_3("3.5.2. Kiểm thử chức năng nghiệp vụ (Functional Testing)")
add_custom_para(
    "Kiểm thử chức năng được tiến hành trực tiếp trên trình duyệt Web để đảm bảo toàn bộ luồng nghiệp vụ liên hoàn hoạt động trơn tru:"
)
test_fn_cols = ["Chức năng kiểm thử", "Các bước thực hiện", "Kết quả mong đợi", "Trạng thái thực tế"]
test_fn_data = [
    ["Luồng Đăng ký & OTP", "1. Nhập form đăng ký\n2. Nhận mail OTP\n3. Điền OTP kích hoạt", "Tài khoản được kích hoạt thành công, đăng nhập bình thường", "ĐẠT (Pass)"],
    ["Luồng Đặt vé & Ví ảo", "1. Chọn gói dịch vụ\n2. Bấm thanh toán bằng số dư ví\n3. Xem chi tiết vé và mã QR", "Ví trừ đúng số tiền, vé tạo trạng thái pending kèm mã QR hợp lệ", "ĐẠT (Pass)"],
    ["Luồng Check-in QR", "1. Đối tác mở Web quét QR\n2. Đưa mã QR vé vào camera\n3. Nhấn check-in", "Mã vé được nhận diện, trạng thái chuyển completed, ghi thời gian check-in", "ĐẠT (Pass)"],
    ["Nhắn tin Realtime", "1. Mở cửa sổ chat với bạn bè\n2. Nhập nội dung và nhấn gửi\n3. Đối phương nhận tin nhắn", "Tin nhắn hiển thị lập tức bên phía đối phương mà không cần reload trang", "ĐẠT (Pass)"],
    ["Duyệt hồ sơ KDL", "1. Admin vào trang quản trị\n2. Chọn hồ sơ KDL\n3. Bấm duyệt (Verify)", "Trạng thái KDL đổi thành verified, bài viết quảng cáo bắt đầu hiển thị trên Feed", "ĐẠT (Pass)"]
]
add_data_dictionary_table("Bảng 3.2: Kết quả kiểm thử các chức năng nghiệp vụ liên hoàn", test_fn_cols, test_fn_data)

add_heading_2("3.6. Đánh giá kết quả hệ thống")

add_heading_3("3.6.1. Ưu điểm")
add_custom_para(
    "1. Kiến trúc phân tầng rõ ràng: Việc triển khai kiến trúc 3 tầng chuẩn kết hợp container Docker giúp hệ thống cực kỳ rõ ràng về mặt cấu trúc mã nguồn, dễ dàng phát triển và bảo trì độc lập giữa Front và Back.\n"
    "2. Nghiệp vụ hoàn chỉnh và khép kín: TravelConnect giải quyết trọn vẹn nghiệp vụ từ khâu chia sẻ kinh nghiệm, tương tác xã hội cho đến đặt vé dịch vụ, thanh toán ví ảo và quét QR Code check-in tại quầy.\n"
    "3. Trải nghiệm thời gian thực cao: Socket.IO mang lại trải nghiệm tương tác trực tiếp (chat, thông báo booking) nhanh chóng và mượt mà.\n"
    "4. Triển khai Cloud hiện đại: Đóng gói Docker Compose giúp ứng dụng dễ dàng triển khai lên bất kỳ đám mây nào một cách nhanh chóng."
)

add_heading_3("3.6.2. Hạn chế")
add_custom_para(
    "1. Cổng thanh toán thực tế: Hệ thống hiện tại mới chỉ tích hợp cổng thanh toán ví ảo mô phỏng nội bộ, chưa kết nối trực tiếp với các cổng thanh toán ngân hàng chính thức như MoMo hay VNPAY.\n"
    "2. Lưu trữ phương tiện: Toàn bộ ảnh tải lên đang được lưu trực tiếp tại ổ đĩa máy chủ Backend thay vì sử dụng các dịch vụ lưu trữ đám mây chuyên dụng như AWS S3, điều này có thể làm đầy ổ đĩa máy chủ khi lượng người dùng tăng cao."
)

add_heading_3("3.6.3. Hướng phát triển tương lai")
add_custom_para(
    "1. Tích hợp thanh toán thật: Kết nối API với cổng thanh toán VNPAY/MoMo để hỗ trợ nạp tiền thật vào ví hệ thống.\n"
    "2. Sử dụng dịch vụ lưu trữ đám mây: Chuyển đổi lưu trữ file upload sang AWS S3 để nâng cao tốc độ tải ảnh và tối ưu hóa tài nguyên máy chủ.\n"
    "3. Phát triển Mobile App: Xây dựng ứng dụng di động đa nền tảng (React Native) để người dùng tiện lợi hơn trong việc check-in QR Code ngay trên điện thoại."
)

doc.add_page_break()

# ----------------- CHAPTER 4 -----------------
add_heading_1("CHƯƠNG 4. TRIỂN KHAI HỆ THỐNG TRÊN CLOUD (CLO3)")

add_heading_2("4.1. Kiến trúc triển khai trên Cloud")

add_heading_3("4.1.1. Sơ đồ kiến trúc triển khai")
add_custom_para(
    "Hệ thống TravelConnect được triển khai thực tế trên môi trường internet toàn cầu sử dụng mô hình kiến trúc Multi-Cloud "
    "tích hợp các dịch vụ đám mây PaaS và SaaS hiện đại. Sơ đồ kiến trúc triển khai thực tế được thiết kế như sau:"
)
add_figure("deployment.png", "Hình 4.1: Sơ đồ kiến trúc triển khai hệ thống trên điện toán đám mây Multi-Cloud")

add_heading_3("4.1.2. Vai trò của các thành phần trong sơ đồ triển khai")
add_custom_para(
    "1. Vercel (Singapore Node): Dịch vụ Cloud PaaS dùng để triển khai Frontend (React + Vite). Tự động tích hợp CI/CD với GitHub, "
    "hỗ trợ SSL bảo mật và CDN giúp tải trang tức thì.\n"
    "2. Render (Singapore Node): Dịch vụ Cloud PaaS dùng để chạy Backend (Node.js Express + Socket.IO). Cung cấp môi trường chạy "
    "Node.js an toàn, tự động cấu hình cổng (port) động và hỗ trợ WebSockets cho các tính năng tương tác thời gian thực.\n"
    "3. Aiven Cloud MySQL: Dịch vụ Cơ sở dữ liệu đám mây SaaS chạy trên nền tảng AWS/GCP Singapore. Đóng vai trò lưu trữ MySQL 8 "
    "an toàn, kết nối bảo mật qua SSL (giao thức bảo mật bắt buộc) ở cổng 24619.\n"
    "4. GitHub Platform: Đóng vai trò là trung tâm quản lý mã nguồn và kích hoạt các đường ống CI/CD tự động. Mỗi khi mã nguồn "
    "có cập nhật (git push), Vercel và Render sẽ tự động kéo code mới nhất về và tự động deploy lại.\n"
    "5. Mock OTP Logs: Hệ thống ghi nhận mã OTP trực tiếp vào nhật ký (Console Logs) của server Render để vượt qua rào cản chặn cổng "
    "SMTP của Render free tier, giúp quá trình đăng ký tài khoản nhanh chóng và bảo mật."
)

add_heading_2("4.2. Quy trình triển khai thực tế trên Cloud")
add_custom_para(
    "Quy trình đưa ứng dụng TravelConnect lên môi trường Internet thực tế được thực hiện qua các bước chi tiết sau:"
)
add_custom_para(
    "Bước 1: Thiết lập Cơ sở dữ liệu MySQL trên Aiven Cloud\n"
    "- Truy cập Aiven.io, tạo mới dịch vụ MySQL bản miễn phí đặt tại Singapore.\n"
    "- Lấy thông tin kết nối an toàn bao gồm Host name, Port (24619), User (avnadmin), Password và SSL Mode (REQUIRED).\n"
    "- Khởi chạy script import để nạp cấu trúc cơ sở dữ liệu và dữ liệu mẫu từ tệp `travelconnect.sql` vào database."
)
add_custom_para(
    "Bước 2: Cấu hình và Deploy Backend trên Render.com\n"
    "- Đăng nhập Render.com qua GitHub, tạo một Web Service kết nối với kho chứa mã nguồn.\n"
    "- Cấu hình thư mục gốc (Root Directory) là `backend`, lệnh Build là `npm install` và lệnh Start là `node server.js`.\n"
    "- Cài đặt các biến môi trường kết nối database Aiven (DB_HOST, DB_PORT, DB_USER, DB_PASS, DB_NAME), khóa bí mật JWT_SECRET, "
    "và đường dẫn CORS của Frontend CLIENT_URL để cho phép kết nối chéo."
)
add_custom_para(
    "Bước 3: Cấu hình và Deploy Frontend trên Vercel.com\n"
    "- Đăng nhập Vercel.com qua GitHub, import repository của dự án.\n"
    "- Chọn thư mục gốc build là `frontend` và cấu hình biến môi trường kết nối API bao gồm: `VITE_API_BASE_URL` (trỏ đến Render Backend "
    "kèm `/api`), `VITE_APP_BASE_URL` và `VITE_SOCKET_URL` (trỏ đến domain Render Backend)."
)
add_custom_para(
    "Bước 4: Cấu hình VerCel SPA Routing\n"
    "- Cấu hình tệp `vercel.json` định tuyến toàn bộ yêu cầu con về `index.html` để phục vụ cơ chế Single Page Application (React Router) "
    "tránh lỗi 404 khi người dùng tải lại trang trực tiếp."
)

add_heading_2("4.3. Minh chứng vận hành trên Cloud")

add_heading_3("4.3.1. Thông tin đường dẫn truy cập")
add_custom_para(
    "Hệ thống hiện tại đã được triển khai chạy thực tế trên Cloud và có thể truy cập trực tiếp thông qua đường dẫn công khai toàn cầu:\n"
    "- Đường dẫn ứng dụng Web chính thức (Frontend): https://travel-connect-chi.vercel.app\n"
    "- Đường dẫn API Endpoint (Backend): https://travelconnect-backend-ovmt.onrender.com/api/health\n"
    "- Đường dẫn cơ sở dữ liệu MySQL (Aiven): mysql-24b8f229-anhlalyn14-32af.j.aivencloud.com:24619"
)

add_heading_3("4.3.2. Hình ảnh minh chứng hệ thống hoạt động thực tế trên Cloud")
add_custom_para(
    "Dưới đây là hình ảnh minh chứng quá trình vận hành thực tế của ứng dụng TravelConnect trên hạ tầng đám mây Vercel và Render:"
)
add_figure("screenshot_homepage.png", "Hình 4.2: Giao diện ứng dụng hoạt động thực tế trên môi trường Vercel Cloud")
add_custom_para(
    "Giao diện chính thức của TravelConnect chạy mượt mà trên tên miền chính thức của Vercel: https://travel-connect-chi.vercel.app. "
    "Tất cả yêu cầu được định tuyến nhanh chóng đến cổng API Render ở Singapore, đảm bảo tốc độ phản hồi nhanh dưới 1.0 giây."
)

doc.add_page_break()

# ----------------- REFERENCES -----------------
add_heading_1("TÀI LIỆU THAM KHẢO")
if orig_doc:
    for idx in range(321, len(orig_doc.paragraphs)):
        p = orig_doc.paragraphs[idx]
        if p.text.strip():
            if "TÀI LIỆU THAM KHẢO" in p.text:
                continue
            add_custom_para(p.text, space_after=Pt(4))
else:
    ref_list = [
        "Buhalis, D., & Law, R. (2008). Progress in information technology and tourism management: 20 years on and 10 years after the Internet—The state of eTourism research. Tourism Management, 29(4), 609–623.",
        "Google. (2023). Google Maps Platform documentation. Retrieved from https://developers.google.com",
        "Nguyễn Văn Minh. (2020). Ứng dụng công nghệ thông tin trong phát triển du lịch thông minh tại Việt Nam. Tạp chí Du lịch Việt Nam, 5(2), 45–50.",
        "PayPal. (2023). PayPal Developer Documentation. Retrieved from https://developer.paypal.com",
        "Phạm Thị Lan. (2021). Chuyển đổi số trong ngành du lịch Việt Nam: Thực trạng và giải pháp. Nhà xuất bản Thống kê.",
        "Pressman, R. S. (2014). Software engineering: A practitioner’s approach (8th ed.). McGraw-Hill.",
        "Sommerville, I. (2016). Software engineering (10th ed.). Pearson Education Limited.",
        "Trần Văn Hùng. (2019). Hệ thống thông tin quản lý. Nhà xuất bản Đại học Kinh tế Quốc dân."
    ]
    for ref in ref_list:
        add_custom_para(ref, space_after=Pt(4))

# Save Document
try:
    doc.save(output_docx)
    print("SUCCESS: Word document report generated successfully at", output_docx)
except PermissionError:
    fallback_docx = os.path.join(base_dir, "docs", "BaoCaoChuyenDe1_PhanDinhLuyen_New.docx")
    doc.save(fallback_docx)
    print("WARNING: Could not overwrite original file because it is open. Saved new version to:", fallback_docx)
